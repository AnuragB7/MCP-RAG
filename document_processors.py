import asyncio
import os
import aiofiles
from typing import Dict, Any, List
import logging
from config import Config

logger = logging.getLogger(__name__)

class ContentExtractor:
    """Extract plain text content from various file formats with large file support"""
    
    @staticmethod
    async def get_file_info(file_path: str) -> Dict[str, Any]:
        """Get file information and determine processing strategy"""
        try:
            file_stat = os.stat(file_path)
            file_size_mb = file_stat.st_size / (1024 * 1024)
            
            return {
                "file_path": file_path,
                "file_size_bytes": file_stat.st_size,
                "file_size_mb": round(file_size_mb, 2),
                "is_large_file": file_size_mb > Config.MEMORY_THRESHOLD_MB,
                "processing_strategy": "chunked" if file_size_mb > Config.MEMORY_THRESHOLD_MB else "standard"
            }
        except Exception as e:
            logger.error(f"Failed to get file info for {file_path}: {e}")
            return {"error": str(e)}
    
    @staticmethod
    async def extract_pdf_content(file_path: str) -> str:
        """Extract all text content from PDF with large file support"""
        try:
            import PyPDF2
            
            file_info = await ContentExtractor.get_file_info(file_path)
            is_large = file_info.get("is_large_file", False)
            
            content = ""
            
            if is_large:
                logger.info(f"Processing large PDF file: {file_path}")
                # Process in chunks for large files
                async with aiofiles.open(file_path, 'rb') as file:
                    file_content = await file.read()
                    
                # Process PDF from memory
                import io
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
                
                # Process pages in batches
                batch_size = 10
                total_pages = len(pdf_reader.pages)
                
                for batch_start in range(0, total_pages, batch_size):
                    batch_end = min(batch_start + batch_size, total_pages)
                    
                    for page_num in range(batch_start, batch_end):
                        page = pdf_reader.pages[page_num]
                        content += page.extract_text() + "\n"
                    
                    # Yield control to prevent blocking
                    await asyncio.sleep(0.01)
                    logger.info(f"Processed pages {batch_start+1}-{batch_end} of {total_pages}")
            else:
                # Standard processing for smaller files
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        content += page.extract_text() + "\n"
            
            logger.info(f"Extracted {len(content)} characters from PDF: {file_path}")
            return content
            
        except Exception as e:
            logger.error(f"Error extracting PDF content from {file_path}: {e}")
            return f"Error extracting PDF content: {str(e)}"
    
    @staticmethod
    async def extract_docx_content(file_path: str) -> str:
        """Extract all text content from DOCX with large file support"""
        try:
            from docx import Document
            
            file_info = await ContentExtractor.get_file_info(file_path)
            is_large = file_info.get("is_large_file", False)
            
            if is_large:
                logger.info(f"Processing large DOCX file: {file_path}")
            
            doc = Document(file_path)
            content = ""
            
            # Extract paragraph text
            paragraph_count = len(doc.paragraphs)
            batch_size = 100 if is_large else paragraph_count
            
            for batch_start in range(0, paragraph_count, batch_size):
                batch_end = min(batch_start + batch_size, paragraph_count)
                
                for i in range(batch_start, batch_end):
                    content += doc.paragraphs[i].text + "\n"
                
                if is_large:
                    await asyncio.sleep(0.01)
                    logger.info(f"Processed paragraphs {batch_start+1}-{batch_end} of {paragraph_count}")
            
            # Extract table text
            if doc.tables:
                logger.info(f"Processing {len(doc.tables)} tables")
                for table_idx, table in enumerate(doc.tables):
                    content += f"\n=== Table {table_idx + 1} ===\n"
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            row_text.append(cell.text.strip())
                        content += " | ".join(row_text) + "\n"
                    
                    if is_large and table_idx % 5 == 0:
                        await asyncio.sleep(0.01)
            
            logger.info(f"Extracted {len(content)} characters from DOCX: {file_path}")
            return content
            
        except Exception as e:
            logger.error(f"Error extracting DOCX content from {file_path}: {e}")
            return f"Error extracting DOCX content: {str(e)}"
    
    @staticmethod
    async def extract_excel_content(file_path: str) -> str:
        """Extract all data from Excel as text with large file support"""
        try:
            import pandas as pd
            
            file_info = await ContentExtractor.get_file_info(file_path)
            is_large = file_info.get("is_large_file", False)
            
            if is_large:
                logger.info(f"Processing large Excel file: {file_path}")
            
            excel_file = pd.ExcelFile(file_path)
            content = ""
            
            for sheet_idx, sheet_name in enumerate(excel_file.sheet_names):
                content += f"\n=== Sheet: {sheet_name} ===\n"
                
                if is_large:
                    # Process large Excel files in chunks
                    chunk_size = 1000
                    chunk_list = []
                    
                    try:
                        for chunk in pd.read_excel(file_path, sheet_name=sheet_name, chunksize=chunk_size):
                            chunk_list.append(chunk)
                            if len(chunk_list) % 5 == 0:
                                await asyncio.sleep(0.01)
                                logger.info(f"Processing chunk {len(chunk_list)} for sheet {sheet_name}")
                        
                        # Combine chunks
                        df = pd.concat(chunk_list, ignore_index=True)
                    except ValueError:
                        # Fallback for sheets that don't support chunking
                        df = pd.read_excel(file_path, sheet_name=sheet_name)
                else:
                    df = pd.read_excel(file_path, sheet_name=sheet_name)
                
                # Convert to string with limited rows for very large sheets
                if len(df) > 10000:
                    content += f"Large sheet with {len(df)} rows, {len(df.columns)} columns\n"
                    content += "First 1000 rows:\n"
                    content += df.head(1000).to_string(max_rows=1000) + "\n"
                    content += f"\n... and {len(df) - 1000} more rows\n"
                else:
                    content += df.to_string() + "\n"
                
                if is_large:
                    await asyncio.sleep(0.01)
                    logger.info(f"Processed sheet {sheet_idx + 1}/{len(excel_file.sheet_names)}: {sheet_name}")
            
            logger.info(f"Extracted {len(content)} characters from Excel: {file_path}")
            return content
            
        except Exception as e:
            logger.error(f"Error extracting Excel content from {file_path}: {e}")
            return f"Error extracting Excel content: {str(e)}"
    
    @staticmethod
    async def extract_csv_content(file_path: str) -> str:
        """Extract all data from CSV as text with large file support"""
        try:
            import pandas as pd
            
            file_info = await ContentExtractor.get_file_info(file_path)
            is_large = file_info.get("is_large_file", False)
            
            if is_large:
                logger.info(f"Processing large CSV file: {file_path}")
                # Process in chunks
                chunk_size = 5000
                chunk_list = []
                
                for chunk in pd.read_csv(file_path, chunksize=chunk_size):
                    chunk_list.append(chunk)
                    if len(chunk_list) % 10 == 0:
                        await asyncio.sleep(0.01)
                        logger.info(f"Processing chunk {len(chunk_list)}")
                
                # Combine chunks
                df = pd.concat(chunk_list, ignore_index=True)
            else:
                df = pd.read_csv(file_path)
            
            # Convert to string with row limits for very large files
            if len(df) > 10000:
                content = f"Large CSV with {len(df)} rows, {len(df.columns)} columns\n"
                content += "Columns: " + ", ".join(df.columns.tolist()) + "\n\n"
                content += "First 1000 rows:\n"
                content += df.head(1000).to_string(max_rows=1000)
                content += f"\n\n... and {len(df) - 1000} more rows"
            else:
                content = df.to_string()
            
            logger.info(f"Extracted {len(content)} characters from CSV: {file_path}")
            return content
            
        except Exception as e:
            logger.error(f"Error extracting CSV content from {file_path}: {e}")
            return f"Error extracting CSV content: {str(e)}"
    
    @staticmethod
    async def extract_content(file_path: str) -> tuple[str, str]:
        """Extract content from any supported file type"""
        file_ext = os.path.splitext(file_path)[1].lower()
        
        extractors = {
            '.pdf': ContentExtractor.extract_pdf_content,
            '.docx': ContentExtractor.extract_docx_content,
            '.xlsx': ContentExtractor.extract_excel_content,
            '.xls': ContentExtractor.extract_excel_content,
            '.csv': ContentExtractor.extract_csv_content
        }
        
        file_type_map = {
            '.pdf': 'PDF',
            '.docx': 'DOCX',
            '.xlsx': 'Excel',
            '.xls': 'Excel',
            '.csv': 'CSV'
        }
        
        if file_ext in extractors:
            content = await extractors[file_ext](file_path)
            file_type = file_type_map[file_ext]
            return content, file_type
        else:
            raise ValueError(f"Unsupported file type: {file_ext}")